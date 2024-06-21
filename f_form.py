import json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from fuzzywuzzy import fuzz

def set_font_style(run, name='Arial', size=12, bold=False):
    run.font.name = name
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold

def format_skills_text(text):
    lines = text.split('\n')
    formatted_lines = []
    for line in lines:
        if line.startswith('• ') and ':' in line:
            category_name_end = line.find(':')
            if category_name_end != -1:
                category_name = line[2:category_name_end].strip()
                formatted_line = f'• **{category_name}**: {line[category_name_end+1:].strip()}'
                formatted_lines.append(formatted_line)
        else:
            formatted_lines.append(line)
    return '\n'.join(formatted_lines)

def format_experience_text(text):
    lines = text.split('\n')
    formatted_lines = []
    for line in lines:
        if line.startswith('*'):
            formatted_lines.append('• ' + line[1:].strip())
        else:
            formatted_lines.append(line)
    return '\n'.join(formatted_lines)

def format_document(profile_path, input_doc_path, output_doc_path):
    # Load the JSON profile file
    with open(profile_path, 'r') as file:
        profile = json.load(file)

    # Load the document
    document = Document(input_doc_path)

    # Extract job titles from JSON profile
    job_titles = [exp["job_title"] for exp in profile["experience"]]
    print("Job Titles from JSON:", job_titles)

    # Extract and format the Skills section
    skills_text = ""
    skills_heading_set = False
    for paragraph in document.paragraphs:
        if 'Skills & Abilities' in paragraph.text:
            skills_heading_set = True
        if skills_heading_set:
            if paragraph.text.startswith('• '):
                skills_text += paragraph.text + '\n'
            if 'Experience' in paragraph.text:
                break

    # Format the extracted Skills text
    formatted_skills_text = format_skills_text(skills_text)

    # Replace the original Skills section with formatted text
    skills_heading_set = False
    for paragraph in document.paragraphs:
        if 'Skills & Abilities' in paragraph.text:
            skills_heading_set = True
        if skills_heading_set:
            if paragraph.text.startswith('• '):
                paragraph.text = formatted_skills_text
                break
            if 'Experience' in paragraph.text:
                break

    # Apply formatting to text between **
    for paragraph in document.paragraphs:
        if '**' in paragraph.text:
            runs = paragraph.runs
            for run in runs:
                if '**' in run.text:
                    parts = run.text.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 1:
                            # This part is between **
                            bold_run = paragraph.add_run(part)
                            set_font_style(bold_run, bold=True)
                        else:
                            normal_run = paragraph.add_run(part)
                            set_font_style(normal_run, bold=False)
                    # Remove the original run
                    paragraph._element.remove(run._element)

    # Adjust Experience section using job titles from JSON profile with fuzzy matching
    experience_text = ""
    experience_heading_set = False
    for paragraph in document.paragraphs:
        if 'Work Experience' in paragraph.text:
            experience_heading_set = True
            print("Found Experience section.")
        if experience_heading_set:
            if paragraph.text.startswith('*'):
                experience_text += paragraph.text + '\n'
            if 'Projects' in paragraph.text:
                break

    # Format the extracted Experience text
    formatted_experience_text = format_experience_text(experience_text)

    # Replace the original Experience section with formatted text
    experience_heading_set = False
    for paragraph in document.paragraphs:
        if 'Work Experience' in paragraph.text:
            experience_heading_set = True
        if experience_heading_set:
            if paragraph.text.startswith('*'):
                paragraph.text = formatted_experience_text
                break
            if 'Projects' in paragraph.text:
                break

    # Apply fuzzy matching to bold job titles in the Experience section
    experience_heading_set = False
    for paragraph in document.paragraphs:
        if 'Work Experience' in paragraph.text:
            experience_heading_set = True
        if experience_heading_set:
            for run in paragraph.runs:
                for job_title in job_titles:
                    if fuzz.partial_ratio(job_title.lower(), run.text.lower()) > 80:  # Adjust the threshold as needed
                        set_font_style(run, bold=True)
                        print("Formatting job title:", run.text)
                        break
            if 'Projects' in paragraph.text:
                break

    # Save the modified document
    document.save(output_doc_path)
    print(f"Formatted document saved to {output_doc_path}")
