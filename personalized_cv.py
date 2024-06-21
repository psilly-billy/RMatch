import os
import json
import streamlit as st
import google.generativeai as genai
from docx import Document
from part_1 import process_template, extract_formatting, replace_text
from generate_summary import generate_summary_section
from generate_skills import generate_skills_section
from generate_experience import generate_experience_section
from f_form import format_document  # Importing the new formatting function

# Load the profile and job listing
def load_json_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return json.load(file)

def configure_gemini(api_key):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')
    return model.start_chat()

def LLM_Response(chat, question):
    response = chat.send_message(question, stream=True)
    response.resolve()  # Ensure the response is fully generated
    return response

def generate_filename(profile_name, job_name, suffix, ext):
    return f"{profile_name[:10]}_{job_name[:10]}_{suffix}.{ext}"

def save_cv_as_word(doc, filename):
    doc.save(f"Data/Gen/{filename}")

def show_personalized_cv():
    st.markdown("<h1>Create Personalized CV</h1>", unsafe_allow_html=True)
    
    st.sidebar.header("Configuration")
    api_key = st.sidebar.text_input("Enter your Gemini API Key", type="password")
    
    profile_files = [f for f in os.listdir("Data/user_profiles/") if f.endswith('.json')]
    job_desc_files = [f for f in os.listdir("Data/JobDescription/") if f.endswith('.json')]
    
    selected_profile = st.selectbox("Select Profile", profile_files)
    selected_job_desc = st.selectbox("Select Job Description", job_desc_files)
    
    if not api_key:
        st.error("Please enter your Gemini API Key.")
        return
    
    if selected_profile and selected_job_desc:
        profile = load_json_file(os.path.join("Data/user_profiles", selected_profile))
        job_listing = load_json_file(os.path.join("Data/JobDescription", selected_job_desc))
        analysis_filename = generate_filename(os.path.splitext(selected_profile)[0], os.path.splitext(selected_job_desc)[0], "analysis", "json")
        analysis_path = os.path.join("Data/analysis/", analysis_filename)
        analysis = load_json_file(analysis_path)
        
        # Process the template to replace placeholders with profile data
        template_path = "Data/ATS_template.docx"
        populated_template_filename = generate_filename(os.path.splitext(selected_profile)[0], os.path.splitext(selected_job_desc)[0], "Populated_ATS_template_v2", "docx")
        populated_template_path = os.path.join("Data/Gen/", populated_template_filename)
        process_template(template_path, os.path.join("Data/user_profiles", selected_profile), populated_template_path)
        
        st.session_state.setdefault('summary', "")
        st.session_state.setdefault('skills', "")
        st.session_state.setdefault('experience', "")

        # Generate and display Summary section
        st.markdown("### Summary Section")
        if st.button("Generate Summary"):
            st.session_state.summary = generate_summary_section(api_key, profile, job_listing, analysis)
        
        if st.session_state.summary:
            st.markdown(f"**Summary:**")
            st.markdown(st.session_state.summary, unsafe_allow_html=True)
            
            if st.button("Save Summary to CV"):
                try:
                    doc = Document(populated_template_path)
                except Exception as e:
                    st.error(f"Failed to load populated template: {e}")
                    return
                
                replacements = {
                    '*Here Summary': st.session_state.summary,
                }
                
                formatting_map = {}
                for para in doc.paragraphs:
                    for run in para.runs:
                        if '*Here Summary' in run.text:
                            formatting_map['*Here Summary'] = extract_formatting(run)
                
                replace_text(doc, replacements, formatting_map)
                save_cv_as_word(doc, populated_template_filename)
                
                st.success("Summary section saved in the CV.")

        # Generate and display Skills section
        #st.markdown("### Skills Section")
        if st.button("Generate Skills"):
            st.session_state.skills = generate_skills_section(api_key, profile, job_listing, analysis)
        
        if st.session_state.skills:
            st.markdown(f"**Skills:**")
            st.markdown(st.session_state.skills, unsafe_allow_html=True)
            
            if st.button("Save Skills to CV"):
                try:
                    doc = Document(populated_template_path)
                except Exception as e:
                    st.error(f"Failed to load populated template: {e}")
                    return
                
                cleaned_skills = st.session_state.skills  # Clean the skills text
                
                replacements = {
                    '*Here Skills': cleaned_skills,
                    '*Here Summary': st.session_state.summary,  # Keep previously saved sections
                }
                
                formatting_map = {}
                for para in doc.paragraphs:
                    for run in para.runs:
                        if '*Here Skills' in run.text:
                            formatting_map['*Here Skills'] = extract_formatting(run)
                        elif '*Here Summary' in run.text:
                            formatting_map['*Here Summary'] = extract_formatting(run)
                
                replace_text(doc, replacements, formatting_map)
                save_cv_as_word(doc, populated_template_filename)
                
                # Format the document with new functions for Skills
                format_document(os.path.join("Data/user_profiles", selected_profile), populated_template_path, populated_template_path)
                
                st.success("Skills section saved in the CV.")

        # Generate and display Experience section
        #st.markdown("### Experience Section")
        if st.button("Generate Experience"):
            st.session_state.experience = generate_experience_section(api_key, profile, job_listing, analysis)
        
        if st.session_state.experience:
            st.markdown(f"**Experience:**")
            st.markdown(st.session_state.experience, unsafe_allow_html=True)
            
            if st.button("Save Experience to CV"):
                try:
                    doc = Document(populated_template_path)
                except Exception as e:
                    st.error(f"Failed to load populated template: {e}")
                    return
                
                replacements = {
                    '*Here Experience': st.session_state.experience,
                    '*Here Skills': st.session_state.skills,      # Keep previously saved sections
                    '*Here Summary': st.session_state.summary,    # Keep previously saved sections
                }
                
                formatting_map = {}
                for para in doc.paragraphs:
                    for run in para.runs:
                        if '*Here Experience' in run.text:
                            formatting_map['*Here Experience'] = extract_formatting(run)
                        elif '*Here Skills' in run.text:
                            formatting_map['*Here Skills'] = extract_formatting(run)
                        elif '*Here Summary' in run.text:
                            formatting_map['*Here Summary'] = extract_formatting(run)
                
                replace_text(doc, replacements, formatting_map)
                save_cv_as_word(doc, populated_template_filename)
                
                # Format the document with new functions for Skills and Experience
                format_document(os.path.join("Data/user_profiles", selected_profile), populated_template_path, populated_template_path)
                
                st.success("Experience section saved in the CV.")
                
            st.markdown("### Download the final CV")
            with open(populated_template_path, "rb") as file:
                st.download_button(
                    label="Download Word Document",
                    data=file,
                    file_name=populated_template_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    show_personalized_cv()
