import os
import json
import streamlit as st
import google.generativeai as genai
from docx import Document
from generate_cl import generate_cover_letter_section

# Load the profile and job listing
def load_json_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return json.load(file)

def configure_gemini(api_key):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')
    return model.start_chat()

def generate_filename(profile_name, job_name, suffix, ext):
    return f"{profile_name[:10]}_{job_name[:10]}_{suffix}.{ext}"

def save_cover_letter_as_word(doc, filename):
    doc.save(f"Data/Gen/{filename}")

def show_cover_letter():
    st.markdown("<h1>Create Personalized Cover Letter</h1>", unsafe_allow_html=True)
    
    st.sidebar.header("Configuration")
    api_key = st.sidebar.text_input("Enter your Gemini API Key", type="password")
    
    profile_files = [f for f in os.listdir("Data/user_profiles/") if f.endswith('.json')]
    job_desc_files = [f for f in os.listdir("Data/JobDescription/") if f.endswith('.json')]
    
    selected_profile = st.selectbox("Select Profile", profile_files)
    selected_job_desc = st.selectbox("Select Job Description", job_desc_files)
    
    if not api_key:
        st.error("Please enter your Gemini API Key.")
        return
    
    if selected_profile and selected_job_desc:
        profile = load_json_file(os.path.join("Data/user_profiles", selected_profile))
        job_listing = load_json_file(os.path.join("Data/JobDescription", selected_job_desc))
        analysis_filename = generate_filename(os.path.splitext(selected_profile)[0], os.path.splitext(selected_job_desc)[0], "analysis", "json")
        analysis_path = os.path.join("Data/analysis/", analysis_filename)
        analysis = load_json_file(analysis_path)
        
        chat = configure_gemini(api_key)
        
        # Generate and display Cover Letter section
        st.markdown("### Cover Letter")
        if st.button("Generate Cover Letter"):
            st.session_state.cover_letter = generate_cover_letter_section(chat, profile, job_listing, analysis)
        
        if st.session_state.get('cover_letter', ""):
            st.markdown(f"**Cover Letter:**")
            st.markdown(st.session_state.cover_letter, unsafe_allow_html=True)
            
            if st.button("Save Cover Letter to File"):
                try:
                    doc = Document()
                    #doc.add_heading('Cover Letter', level=1)
                    doc.add_paragraph(st.session_state.cover_letter)
                    
                    profile_name = os.path.splitext(selected_profile)[0]
                    job_name = os.path.splitext(selected_job_desc)[0]
                    word_filename = generate_filename(profile_name, job_name, "cover_letter", "docx")
                    
                    save_cover_letter_as_word(doc, word_filename)
                    
                    st.success("Cover letter generated and saved successfully.")
                    st.markdown(f"[Download Word Document](./Data/Gen/{word_filename})")
                    
                    with open(f"Data/Gen/{word_filename}", "rb") as word_file:
                        st.download_button(
                            label="Download Word Document",
                            data=word_file,
                            file_name=f"{profile_name[:10]}_{job_name[:10]}_cover_letter.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Failed to save cover letter: {e}")

if __name__ == "__main__":
    show_cover_letter()
