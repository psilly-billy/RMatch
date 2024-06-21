from docx import Document
import docx.opc.constants
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import json
import os


def extract_formatting(run):
    """Extract formatting from a run."""
    if run is None:
        return {
            'font_name': 'Arial',
            'font_size': 12,
            'bold': False,
            'italic': False,
            'underline': False,
            'color': None
        }
    return {
        'font_name': run.font.name,
        'font_size': run.font.size.pt if run.font.size else 11,
        'bold': run.font.bold,
        'italic': run.font.italic,
        'underline': run.font.underline,
        'color': run.font.color.rgb if run.font.color else None
    }

def apply_formatting(run, formatting):
    """Apply formatting to a run."""
    run.font.name = formatting['font_name']
    run.font.size = Pt(formatting['font_size'])
    run.font.bold = formatting['bold']
    run.font.italic = formatting['italic']
    run.font.underline = formatting['underline']
    if formatting['color']:
        run.font.color.rgb = formatting['color']

def add_hyperlink(paragraph, url, text, formatting):
    """
    A function that places a hyperlink within a paragraph object.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    
    r_style = OxmlElement('w:rStyle')
    r_style.set(qn('w:val'), 'Hyperlink')
    r_pr.append(r_style)
    
    # Font settings
    r_fonts = OxmlElement('w:rFonts')
    font_name = formatting['font_name'] or 'Arial'
    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)
    r_pr.append(r_fonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(formatting['font_size'] * 2)))
    r_pr.append(sz)
    
    new_run.append(r_pr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._element.append(hyperlink)

def replace_text_in_element(element, replacements, formatting_map):
    for key, value in replacements.items():
        if key in element.text:
            if key == '*Linkedin' or key == '*Github':
                if element.runs:
                    formatting = formatting_map.get(key, extract_formatting(element.runs[0]))
                else:
                    formatting = {'font_name': 'Arial', 'font_size': 12, 'bold': False, 'italic': False, 'underline': False, 'color': None}
                element.clear()
                add_hyperlink(element, value, value, formatting)
            else:
                for run in element.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        apply_formatting(run, formatting_map.get(key, extract_formatting(run)))
                element.text = element.text.replace(key, value)

def replace_name(doc, name, formatting_map):
    for para in doc.paragraphs:
        for run in para.runs:
            if '*Name' in run.text:
                run.text = run.text.replace('*Name', name)
                apply_formatting(run, formatting_map.get('*Name', extract_formatting(run)))

def replace_text(doc, replacements, formatting_map):
    for para in doc.paragraphs:
        replace_text_in_element(para, replacements, formatting_map)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_element(para, replacements, formatting_map)

def replace_projects_and_education(doc, data, formatting_map):
    # Replace projects
    for para in doc.paragraphs:
        if '*Here Projects' in para.text:
            para.clear()
            first_project = True
            for proj in data.get('projects', []):
                if not first_project:
                    para.add_run('\n')
                first_project = False
                formatting = formatting_map.get('*Here Projects', extract_formatting(para.runs[0] if para.runs else None))
                add_hyperlink(para, proj['url'], proj['title'], formatting)
                proj_run = para.add_run('\n' + proj['description'] + '\n')
                apply_formatting(proj_run, formatting)

    # Replace education and certifications
    for para in doc.paragraphs:
        if '*Here Education' in para.text:
            para.clear()
            first_entry = True
            for edu in data.get('education', []):
                if not first_entry:
                    para.add_run('\n')
                first_entry = False
                run = para.add_run(f"{edu['degree']} from {edu['institution']} ({edu['location']})\n{edu['start_date']} - {edu['end_date']}\n")
                apply_formatting(run, formatting_map.get('*Here Education', extract_formatting(run)))
            for cert in data.get('certifications', []):
                if not first_entry:
                    para.add_run('\n')
                first_entry = False
                formatting = formatting_map.get('*Here Education', extract_formatting(para.runs[0] if para.runs else None))
                add_hyperlink(para, cert['description'], cert['certification'], formatting)
                cert_run = para.add_run(f"\n{cert['certification']} by {cert['certifying_body']} on {cert['date_awarded']}\n")
                apply_formatting(cert_run, formatting)

# Detailed check for all elements
def check_and_replace(doc, replacements, formatting_map):
    for paragraph in doc.paragraphs:
        replace_text_in_element(paragraph, replacements, formatting_map)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_element(paragraph, replacements, formatting_map)

# Main function to process the template
def process_template(template_path, profile_path, output_path):
    # Load the template document
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Failed to load template: {e}")
        return

    # Load the JSON data
    try:
        with open(profile_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"Failed to load JSON data: {e}")
        return

    # Replace placeholders with actual data
    replacements = {
        '*Name': data.get('name', ''),
        '*Location': data.get('location', ''),
        '*Email': data.get('email', ''),
        '*Phone Number': data.get('phone', ''),
        '*Linkedin': data.get('linkedin', ''),
        '*Github': data.get('github', ''),
    }

    # Create a map of formatting settings from the template
    formatting_map = {}
    for para in doc.paragraphs:
        for run in para.runs:
            if '*Name' in run.text:
                formatting_map['*Name'] = extract_formatting(run)
            elif '*Linkedin' in run.text:
                formatting_map['*Linkedin'] = extract_formatting(run)
            elif '*Github' in run.text:
                formatting_map['*Github'] = extract_formatting(run)
            elif '*Here Projects' in run.text:
                formatting_map['*Here Projects'] = extract_formatting(run)
            elif '*Here Education' in run.text:
                formatting_map['*Here Education'] = extract_formatting(run)

    # Replace name separately to ensure it matches the template's formatting
    replace_name(doc, replacements['*Name'], formatting_map)
    replace_text(doc, replacements, formatting_map)
    replace_projects_and_education(doc, data, formatting_map)
    check_and_replace(doc, replacements, formatting_map)


    
    #output_path = (f"Data/Gen/{output_path}")

    # Save the modified document
    try:
        doc.save(output_path)
        print(f"Document saved as {output_path}")
    except Exception as e:
        print(f"Failed to save document: {e}")

    # Verify the file was saved correctly
    if os.path.exists(output_path):
        print(f"File saved successfully at {output_path}")
    else:
        print("Error: File was not saved.")


