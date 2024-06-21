import os
import json
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor


def generate_filename(profile_name, job_name, suffix, ext):
    return f"{profile_name[:10]}_{job_name[:10]}_{suffix}.{ext}"

def extract_formatting(run):
    if run is None:
        return {
            'font_name': 'Arial',
            'font_size': 11,
            'bold': False,
            'italic': False,
            'underline': False,
            'color': None
        }
    return {
        'font_name': run.font.name,
        'font_size': run.font.size.pt if run.font.size else 11,
        'bold': run.font.bold,
        'italic': run.font.italic,
        'underline': run.font.underline,
        'color': run.font.color.rgb if run.font.color else None
    }

def apply_formatting(run, formatting):
    run.font.name = formatting['font_name']
    run.font.size = Pt(formatting['font_size'])
    run.bold = formatting['bold']
    run.italic = formatting['italic']
    run.underline = formatting['underline']
    if formatting['color']:
        run.font.color.rgb = formatting['color']

def format_section_titles(paragraph):
   for run in paragraph.runs:
        text = run.text
        if ':' in text:
            before_colon, after_colon = text.split(':', 1)
            run.text = before_colon + ':'
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

            new_run = paragraph.add_run(after_colon)
            new_run.font.name = 'Arial'
            new_run.font.size = Pt(11)
            new_run.bold = False
            return

def replace_text_in_element(element, replacements, formatting_map):
    for key, value in replacements.items():
        if key in element.text:
            if key in formatting_map:
                run_formatting = formatting_map[key]
            else:
                run_formatting = extract_formatting(None)
            for run in element.runs:
                run.clear()
                run.text = value
                apply_formatting(run, run_formatting)

def update_section(doc, section_key, new_text, formatting_map):
    replacements = {section_key: new_text}
    for para in doc.paragraphs:
        replace_text_in_element(para, replacements, formatting_map)
        if section_key in para.text:
            format_section_titles(para)

def clean_skills_text(skills_text):
    cleaned_text = skills_text.replace('**', '').replace('## Skills & Abilities\n\n', '')
    return cleaned_text

def clean_experience_text(experience_text):
    cleaned_text = experience_text.replace('**', '').replace('* ', '• ').replace('**Work Experience**', '')
    return cleaned_text


def save_cv_as_word(doc, filename):
    doc.save(f"Data/Gen/{filename}")

def load_saved_doc(populated_template_path):
    try:
        return Document(populated_template_path)
    except Exception as e:
        st.error(f"Failed to load populated template: {e}")
        return None